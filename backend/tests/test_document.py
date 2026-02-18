"""Unit tests for app.services.document — type detection, PDF/DOCX extraction."""

import io
import pytest

from app.services.document import (
    detect_document_type,
    extract_pdf,
    extract_docx,
    DocumentResult,
)


# ---------------------------------------------------------------------------
# detect_document_type
# ---------------------------------------------------------------------------


class TestDetectDocumentType:

    def test_pdf_by_url_extension(self):
        assert detect_document_type("https://example.com/report.pdf") == "pdf"

    def test_pdf_by_url_extension_case_insensitive(self):
        assert detect_document_type("https://example.com/report.PDF") == "pdf"

    def test_pdf_by_url_with_query_params(self):
        assert detect_document_type("https://example.com/file.pdf?token=abc") == "pdf"

    def test_pdf_by_url_with_fragment(self):
        assert detect_document_type("https://example.com/file.pdf#page=2") == "pdf"

    def test_docx_by_url_extension(self):
        assert detect_document_type("https://example.com/doc.docx") == "docx"

    def test_doc_by_url_extension(self):
        """A .doc extension maps to 'docx' (attempted docx parsing)."""
        assert detect_document_type("https://example.com/old.doc") == "docx"

    def test_xlsx_by_url_extension(self):
        assert detect_document_type("https://example.com/data.xlsx") == "xlsx"

    def test_pdf_by_content_type(self):
        assert detect_document_type(
            "https://example.com/download",
            content_type="application/pdf",
        ) == "pdf"

    def test_docx_by_content_type(self):
        ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert detect_document_type("https://example.com/dl", content_type=ct) == "docx"

    def test_xlsx_by_content_type(self):
        ct = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        assert detect_document_type("https://example.com/dl", content_type=ct) == "xlsx"

    def test_msword_content_type(self):
        assert detect_document_type(
            "https://example.com/dl",
            content_type="application/msword",
        ) == "docx"

    def test_html_by_content_type(self):
        assert detect_document_type(
            "https://example.com/page",
            content_type="text/html; charset=utf-8",
        ) == "html"

    def test_xhtml_by_content_type(self):
        assert detect_document_type(
            "https://example.com/page",
            content_type="application/xhtml+xml",
        ) == "html"

    def test_pdf_by_magic_bytes(self):
        assert detect_document_type(
            "https://example.com/download",
            raw_bytes=b"%PDF-1.4 fake pdf content",
        ) == "pdf"

    def test_zip_based_defaults_to_docx(self):
        """PK ZIP magic without .xlsx extension defaults to docx."""
        assert detect_document_type(
            "https://example.com/unknown",
            raw_bytes=b"PK\x03\x04some-zip-content",
        ) == "docx"

    def test_zip_based_with_xlsx_extension(self):
        """PK ZIP magic + .xlsx extension returns xlsx."""
        assert detect_document_type(
            "https://example.com/file.xlsx",
            raw_bytes=b"PK\x03\x04some-zip-content",
        ) == "xlsx"

    def test_defaults_to_html(self):
        """When nothing else matches, default is html."""
        assert detect_document_type("https://example.com/page") == "html"

    def test_content_type_overrides_when_no_extension(self):
        """Content-type takes precedence when URL has no extension."""
        assert detect_document_type(
            "https://example.com/api/file/12345",
            content_type="application/pdf",
        ) == "pdf"

    def test_url_extension_takes_precedence_over_content_type(self):
        """URL extension is checked before content-type."""
        result = detect_document_type(
            "https://example.com/file.pdf",
            content_type="text/html",
        )
        assert result == "pdf"


# ---------------------------------------------------------------------------
# extract_pdf (uses PyMuPDF / fitz)
# ---------------------------------------------------------------------------


def _create_minimal_pdf(text: str = "Hello, WebHarvest!", title: str = "Test PDF") -> bytes:
    """Programmatically create a minimal PDF using PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text, fontsize=12)
    doc.set_metadata({"title": title, "author": "Test Suite"})
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


class TestExtractPdf:

    @pytest.mark.asyncio
    async def test_extracts_text_from_pdf(self):
        """Text inserted into a PDF is extracted correctly."""
        pdf_bytes = _create_minimal_pdf("Hello, WebHarvest!")
        result = await extract_pdf(pdf_bytes)

        assert isinstance(result, DocumentResult)
        assert "Hello" in result.text
        assert "WebHarvest" in result.text

    @pytest.mark.asyncio
    async def test_page_count(self):
        """Page count is correct for a single-page PDF."""
        pdf_bytes = _create_minimal_pdf("Page one")
        result = await extract_pdf(pdf_bytes)
        assert result.page_count == 1

    @pytest.mark.asyncio
    async def test_word_count(self):
        """Word count reflects the text content."""
        pdf_bytes = _create_minimal_pdf("one two three four five")
        result = await extract_pdf(pdf_bytes)
        assert result.word_count >= 5

    @pytest.mark.asyncio
    async def test_metadata_extraction(self):
        """PDF metadata (title, author) is extracted."""
        pdf_bytes = _create_minimal_pdf("Content", title="My Report")
        result = await extract_pdf(pdf_bytes)
        assert result.metadata.get("title") == "My Report"
        assert result.metadata.get("author") == "Test Suite"
        assert result.metadata.get("document_type") == "pdf"

    @pytest.mark.asyncio
    async def test_markdown_output(self):
        """Markdown output includes page headings and text."""
        pdf_bytes = _create_minimal_pdf("Important content here")
        result = await extract_pdf(pdf_bytes)
        assert "## Page 1" in result.markdown
        assert "Important content here" in result.markdown

    @pytest.mark.asyncio
    async def test_multi_page_pdf(self):
        """A multi-page PDF reports correct page count."""
        import fitz

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content", fontsize=12)
        pdf_bytes = doc.tobytes()
        doc.close()

        result = await extract_pdf(pdf_bytes)
        assert result.page_count == 3
        assert "Page 1 content" in result.text
        assert "Page 3 content" in result.text

    @pytest.mark.asyncio
    async def test_corrupted_pdf_returns_error_result(self):
        """Corrupted PDF data returns an error DocumentResult."""
        result = await extract_pdf(b"not-a-real-pdf-at-all")
        assert "failed" in result.text.lower() or "error" in result.metadata.get("error", "").lower()


# ---------------------------------------------------------------------------
# extract_docx (uses python-docx)
# ---------------------------------------------------------------------------


def _create_minimal_docx(
    paragraphs: list[str] | None = None,
    title: str = "Test Document",
    author: str = "Test Suite",
) -> bytes:
    """Programmatically create a minimal DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author

    if paragraphs is None:
        paragraphs = ["Hello from DOCX!", "This is paragraph two."]

    for text in paragraphs:
        doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractDocx:

    @pytest.mark.asyncio
    async def test_extracts_text_from_docx(self):
        """Text from paragraphs is extracted."""
        docx_bytes = _create_minimal_docx(["Alpha", "Beta", "Gamma"])
        result = await extract_docx(docx_bytes)

        assert isinstance(result, DocumentResult)
        assert "Alpha" in result.text
        assert "Beta" in result.text
        assert "Gamma" in result.text

    @pytest.mark.asyncio
    async def test_metadata_extraction(self):
        """DOCX metadata (title, author) is extracted."""
        docx_bytes = _create_minimal_docx(title="Quarterly Report", author="Jane Doe")
        result = await extract_docx(docx_bytes)
        assert result.metadata.get("title") == "Quarterly Report"
        assert result.metadata.get("author") == "Jane Doe"
        assert result.metadata.get("document_type") == "docx"

    @pytest.mark.asyncio
    async def test_word_count(self):
        """Word count reflects paragraph content."""
        docx_bytes = _create_minimal_docx(["one two three", "four five six seven"])
        result = await extract_docx(docx_bytes)
        assert result.word_count >= 7

    @pytest.mark.asyncio
    async def test_markdown_output(self):
        """Markdown output includes the document title and content."""
        docx_bytes = _create_minimal_docx(
            paragraphs=["Important finding"],
            title="Research Paper",
        )
        result = await extract_docx(docx_bytes)
        assert "Research Paper" in result.markdown
        assert "Important finding" in result.markdown

    @pytest.mark.asyncio
    async def test_heading_conversion(self):
        """Heading styles are converted to markdown headings."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.core_properties.title = "Test"
        doc.add_heading("Main Heading", level=1)
        doc.add_paragraph("Body text")
        doc.add_heading("Sub Heading", level=2)

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = await extract_docx(docx_bytes)
        assert "# Main Heading" in result.markdown
        assert "## Sub Heading" in result.markdown

    @pytest.mark.asyncio
    async def test_table_extraction(self):
        """Tables in DOCX are included in the output."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.core_properties.title = "Test"
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"
        table.cell(1, 0).text = "Val1"
        table.cell(1, 1).text = "Val2"

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = await extract_docx(docx_bytes)
        assert "Header1" in result.text or "Header1" in result.markdown
        assert "Val1" in result.text or "Val1" in result.markdown
        assert result.metadata.get("table_count", 0) >= 1

    @pytest.mark.asyncio
    async def test_corrupted_docx_returns_error_result(self):
        """Invalid DOCX bytes return an error DocumentResult."""
        result = await extract_docx(b"not-a-real-docx")
        assert "failed" in result.text.lower() or "error" in result.metadata.get("error", "").lower()

    @pytest.mark.asyncio
    async def test_page_count_is_one(self):
        """DOCX extraction reports page_count=1 (no fixed pages in docx)."""
        docx_bytes = _create_minimal_docx()
        result = await extract_docx(docx_bytes)
        assert result.page_count == 1

    @pytest.mark.asyncio
    async def test_paragraph_count_metadata(self):
        """Metadata includes a paragraph_count field."""
        docx_bytes = _create_minimal_docx(["Para 1", "Para 2", "Para 3"])
        result = await extract_docx(docx_bytes)
        assert result.metadata.get("paragraph_count", 0) >= 3
