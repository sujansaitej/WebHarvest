"""Document extraction service for PDF, DOCX, and other non-HTML formats."""

import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class DocumentResult:
    """Extracted document content."""
    text: str = ""
    markdown: str = ""
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    word_count: int = 0


def detect_document_type(
    url: str,
    content_type: str | None = None,
    raw_bytes: bytes = b"",
) -> str:
    """Detect document type from URL extension and content-type header.

    Returns: "html", "pdf", "docx", "xlsx", or "unknown"
    """
    url_lower = url.lower().split("?")[0].split("#")[0]

    # Check URL extension first
    if url_lower.endswith(".pdf"):
        return "pdf"
    if url_lower.endswith(".docx"):
        return "docx"
    if url_lower.endswith(".xlsx"):
        return "xlsx"
    if url_lower.endswith(".doc"):
        return "docx"  # Will attempt docx parsing

    # Check content-type header
    if content_type:
        ct = content_type.lower()
        if "application/pdf" in ct:
            return "pdf"
        if "application/vnd.openxmlformats-officedocument.wordprocessingml" in ct:
            return "docx"
        if "application/vnd.openxmlformats-officedocument.spreadsheetml" in ct:
            return "xlsx"
        if "application/msword" in ct:
            return "docx"
        if "text/html" in ct or "application/xhtml" in ct:
            return "html"

    # Check magic bytes
    if raw_bytes:
        if raw_bytes[:4] == b"%PDF":
            return "pdf"
        if raw_bytes[:4] == b"PK\x03\x04":
            # ZIP-based format (docx, xlsx, etc.)
            if url_lower.endswith(".xlsx"):
                return "xlsx"
            return "docx"  # Default to docx for ZIP-based

    return "html"  # Default to HTML


async def extract_pdf(raw_bytes: bytes) -> DocumentResult:
    """Extract text and metadata from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF not installed, returning empty result")
        return DocumentResult(
            text="[PDF extraction requires PyMuPDF]",
            markdown="*PDF extraction requires PyMuPDF*",
            metadata={"error": "PyMuPDF not installed"},
        )

    try:
        doc = fitz.open(stream=raw_bytes, filetype="pdf")

        # Extract metadata
        meta = doc.metadata or {}
        metadata = {
            "author": meta.get("author", ""),
            "title": meta.get("title", ""),
            "subject": meta.get("subject", ""),
            "creator": meta.get("creator", ""),
            "producer": meta.get("producer", ""),
            "creation_date": meta.get("creationDate", ""),
            "mod_date": meta.get("modDate", ""),
            "page_count": doc.page_count,
            "document_type": "pdf",
        }

        # Extract table of contents
        toc = doc.get_toc()
        if toc:
            metadata["table_of_contents"] = [
                {"level": level, "title": title, "page": page}
                for level, title, page in toc
            ]

        # Extract text from all pages
        pages_text = []
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                pages_text.append(text)

        full_text = "\n\n".join(pages_text)
        word_count = len(full_text.split())

        # Build markdown output
        md_parts = []
        title = metadata.get("title") or "PDF Document"
        md_parts.append(f"# {title}\n")
        if metadata.get("author"):
            md_parts.append(f"**Author:** {metadata['author']}\n")
        md_parts.append(f"**Pages:** {doc.page_count} | **Words:** {word_count}\n")
        md_parts.append("---\n")

        for i, page_text in enumerate(pages_text):
            md_parts.append(f"## Page {i + 1}\n")
            md_parts.append(page_text.strip())
            md_parts.append("")

        markdown = "\n\n".join(md_parts)

        doc.close()

        return DocumentResult(
            text=full_text,
            markdown=markdown,
            metadata=metadata,
            page_count=doc.page_count,
            word_count=word_count,
        )

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return DocumentResult(
            text=f"[PDF extraction failed: {e}]",
            markdown=f"*PDF extraction failed: {e}*",
            metadata={"error": str(e), "document_type": "pdf"},
        )


async def extract_docx(raw_bytes: bytes) -> DocumentResult:
    """Extract text and metadata from a DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed, returning empty result")
        return DocumentResult(
            text="[DOCX extraction requires python-docx]",
            markdown="*DOCX extraction requires python-docx*",
            metadata={"error": "python-docx not installed"},
        )

    try:
        doc = Document(io.BytesIO(raw_bytes))

        # Extract metadata
        props = doc.core_properties
        metadata = {
            "author": props.author or "",
            "title": props.title or "",
            "subject": props.subject or "",
            "created": props.created.isoformat() if props.created else "",
            "modified": props.modified.isoformat() if props.modified else "",
            "last_modified_by": props.last_modified_by or "",
            "revision": props.revision,
            "document_type": "docx",
        }

        # Extract paragraphs with heading detection
        md_parts = []
        text_parts = []

        title = metadata.get("title") or "Document"
        md_parts.append(f"# {title}\n")
        if metadata.get("author"):
            md_parts.append(f"**Author:** {metadata['author']}\n")
        md_parts.append("---\n")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text_parts.append(text)

            # Convert heading styles to markdown
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading 1"):
                md_parts.append(f"# {text}")
            elif style_name.startswith("Heading 2"):
                md_parts.append(f"## {text}")
            elif style_name.startswith("Heading 3"):
                md_parts.append(f"### {text}")
            elif style_name.startswith("Heading 4"):
                md_parts.append(f"#### {text}")
            elif style_name == "List Bullet":
                md_parts.append(f"- {text}")
            elif style_name == "List Number":
                md_parts.append(f"1. {text}")
            else:
                md_parts.append(text)

        # Extract tables
        for table in doc.tables:
            md_parts.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                md_parts.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    md_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
                for cell_text in cells:
                    if cell_text:
                        text_parts.append(cell_text)

        full_text = "\n".join(text_parts)
        word_count = len(full_text.split())
        markdown = "\n\n".join(md_parts)

        metadata["word_count"] = word_count
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        return DocumentResult(
            text=full_text,
            markdown=markdown,
            metadata=metadata,
            page_count=1,  # DOCX doesn't have fixed pages
            word_count=word_count,
        )

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return DocumentResult(
            text=f"[DOCX extraction failed: {e}]",
            markdown=f"*DOCX extraction failed: {e}*",
            metadata={"error": str(e), "document_type": "docx"},
        )
